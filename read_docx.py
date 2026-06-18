# read_docx.py
'''
read_docx.py
Все функции для работы с Word-документами (.doc, .docx)
Извлечение таблиц, конвертация, перемещение файлов
'''

import os
import shutil
import re
import win32com.client
from pathlib import Path
from docx import Document
import config

# ========== ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ (общие) ==========

def make_hashable(item):
    """Рекурсивно преобразует списки в кортежи для возможности хеширования.
       Нужна для is_subset_with_lists."""
    if isinstance(item, list):
        return tuple(make_hashable(x) for x in item)
    return item

def is_subset_with_lists(list1, list2):
    """Проверяет, является ли list1 подмножеством list2 с учётом вложенных списков.
       Используется для склейки разорванных строк."""
    set1 = {make_hashable(x) for x in list1}
    set2 = {make_hashable(x) for x in list2}
    return set1.issubset(set2)

def all_deep_empty(lst):
    """Рекурсивно проверяет, пусты ли все вложенные списки."""
    if not isinstance(lst, list):
        return False
    if not lst:
        return True
    return all(all_deep_empty(item) for item in lst)

def get_mismatch_indices(list1, list2):
    """Возвращает индексы элементов, которые различаются в двух списках.
       Списки должны быть одинаковой длины."""
    if len(list1) != len(list2):
        raise ValueError("Списки должны быть одинаковой длины")
    mismatch_indices = []
    for i, (a, b) in enumerate(zip(list1, list2)):
        if a != b:
            mismatch_indices.append(i)
    return mismatch_indices

# ----- блок функций для конвертации и работы с Word (существующие) -----

def convert_doc_to_docx(doc_path):
    '''
    Конвертирует .doc в .docx
    '''
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    doc = None
    try:
        doc = word.Documents.Open(os.path.abspath(doc_path))
        docx_path = doc_path.replace('.doc', '.docx')
        doc.SaveAs(docx_path, FileFormat=16)
        doc.Close()
        word.Quit()
        return docx_path
    except Exception as e:
        if doc:
            doc.Close()
        word.Quit()
        raise

def convert_numbering_to_text(doc_path):
    """Конвертирует автоматическую нумерацию в текст в документе Word"""
    abs_path = os.path.abspath(doc_path)
    word = None
    doc = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False
        doc = word.Documents.Open(abs_path)
        doc.ConvertNumbersToText()
        doc.SaveAs(abs_path)
        return Path(abs_path)
    except Exception as e:
        print(f"Ошибка конвертации нумерации: {e}")
        return None
    finally:
        if doc:
            doc.Close()
        if word:
            word.Quit()

def move_file(source_path, destination_dir):
    """
    Переносит файл из source_path в папку destination_dir
    
    Args:
        source_path: полный путь к файлу (например, "C:/files/doc.docx")
        destination_dir: путь к папке назначения (например, "C:/archive")
    """
    if not os.path.exists(source_path):
        print(f"Файл не найден: {source_path}")
        return False

    try:    
        os.makedirs(destination_dir, exist_ok=True)
        filename = os.path.basename(source_path)
        destination_path = os.path.join(destination_dir, filename)
        shutil.move(source_path, destination_path)
        print(f"Файл перемещён: {source_path} → {destination_path}")
        return True
    except Exception as e:
        print(f"Ошибка при переносе: {e}")
        return False

def take_all_docx_from_dir(journals_directory):
    '''
    получение списка всех .docx файлов из указанной директории
    все .doc файлы преобразуются в .docx и перемещаются в отдельную директорию
    Возвращает: (список_файлов, список_конвертированных)
    '''
    source_dir = Path(journals_directory).resolve()
    
    docx_files = list(source_dir.glob('*.docx')) 
    doc_files = list(source_dir.glob('*.doc'))
    existing_stems = {f.stem for f in docx_files}
    files = docx_files.copy()
    converted = []  # список конвертированных файлов
    
    for doc_file in doc_files:
        if doc_file.stem not in existing_stems:
            new_path = convert_doc_to_docx(str(doc_file))
            files.append(Path(new_path))
            converted.append((doc_file.name, Path(new_path).name))
        try:
            doc_backup_dir = source_dir / 'doc files'
            doc_backup_dir.mkdir(exist_ok=True)
            shutil.move(str(doc_file), str(doc_backup_dir / doc_file.name))
        except Exception as e:
            print(f"  ⚠️ Не удалось переместить {doc_file.name}: {e}")
    return files, converted


# ----- блок функций для извлечения данных из таблиц docx -----

def extract_all_text_from_cell(cell):
    """Рекурсивно извлекает весь текст из ячейки, включая вложенные таблицы"""
    result = []
    if cell:
        for para in cell.paragraphs:
            if para.text.strip():
                text = para.text.strip()
                text = ' '.join(text.split())
                text = text.replace('\xad', '-').replace('\n', ' ')
                result.append(text)
        for inner_table in cell.tables:
            for it_row in inner_table.rows:
                for nested_cell in it_row.cells:
                    inner_table_result = extract_all_text_from_cell(nested_cell)
                    result.extend(inner_table_result)
    return result

def remove_duplicate_pairs(lst):
    """Удаляет дублирующиеся пары элементов в списке (для обработки объединённых ячеек)"""
    if len(lst) % 2 != 0:
        return lst
    result = []
    for i in range(0, len(lst), 2):
        if lst[i] != lst[i + 1]:
            return lst
        result.append(lst[i])
    return result

def remove_duplicates(lst):
    """Удаляет последовательные дубликаты в списке, но сохраняет группы из трёх одинаковых элементов"""
    if not lst:
        return []
    result = []
    skip = 0
    for i in range(len(lst)):
        if skip > 0:
            skip -= 1
            continue
        if i + 2 < len(lst) and lst[i] == ['0'] and lst[i + 1] == ['0'] and lst[i + 2] == ['0']:
            result.extend([['0'], ['0'], ['0']])
            skip = 2
            continue
        if i + 2 < len(lst) and lst[i] == ['-'] and lst[i + 1] == ['-'] and lst[i + 2] == ['-']:
            result.extend([['-'], ['-'], ['-']])
            skip = 2
            continue
        result.append(lst[i])
        j = i + 1
        while j < len(lst) and lst[j] == lst[i]:
            j += 1
        skip = j - i - 1
    return result

def find_first_and_last_sublist_index(big_list, small_list):
    """Находит первый и последний индекс вхождения подсписка в большой список"""
    len_small = len(small_list)
    len_big = len(big_list)
    first_index = -1
    last_index = -1
    for i in range(len_big - len_small + 1):
        if big_list[i:i + len_small] == small_list:
            if first_index == -1:
                first_index = i
            last_index = i
    return first_index, last_index

def extract_all_text_from_table(table):
    """
    Извлекает и очищает данные из таблицы Word.
    Обрабатывает объединённые ячейки, вложенные таблицы, пустые строки,
    удаляет дубликаты, склеивает разорванные строки, парсит координаты.
    """
    if not table.rows:
        return []
    table_contents = []
    try:
        for row in table.rows:
            current_row = []
            for cell in row.cells:
                current_row.append(extract_all_text_from_cell(cell))
            if len(current_row) > 6 and not all_deep_empty(current_row):
                current_row = remove_duplicate_pairs(current_row)
                table_contents.append(current_row)
    except Exception as e:
        print(f'Ошибка! extract_all_text_from_table (0) - {e}')

    if not table_contents:
        return []

    # Чистка пустых столбцов
    new_table_contents = []
    try:
        for row in table_contents:
            current_row = row
            if all_deep_empty(row[0]) and config.regular_num.search(' '.join(row[1])):
                current_row = row[1:]
            if all_deep_empty(row[-1]) and (any(config.regular_letter_minus.search(s) for s in row[-2]) or all_deep_empty(row[-2])):
                current_row = current_row[:-1]
            new_table_contents.append(current_row)
        table_contents = new_table_contents
    except Exception as e:
        print(f'Ошибка! extract_all_text_from_table (1) - {e}')

    # Удаление строк без номера
    try:
        list_for_delete = []
        for index, row in enumerate(table_contents):
            if row and row[0] and row[0][0]:
                match = config.regular_num.search(row[0][0])
                if match:
                    table_contents[index][0] = [match.group()]
                else:
                    list_for_delete.append(index)
            else:
                list_for_delete.append(index)
        table_contents = [item for idx, item in enumerate(table_contents) if idx not in list_for_delete]
    except Exception as e:
        print(f'Ошибка! extract_all_text_from_table (3) - {e}')

    # Удаление дублей из-за вложенных таблиц
    try:
        list_for_delete = []
        for i in range(len(table_contents) - 1):
            if (len(table_contents[i]) == len(table_contents[i + 1])):
                mismatch = get_mismatch_indices(table_contents[i], table_contents[i + 1])
                if mismatch == [2]:
                    list_for_delete.append(i + 1)
        table_contents = [item for idx, item in enumerate(table_contents) if idx not in list_for_delete]
    except Exception as e:
        print(f'Ошибка! extract_all_text_from_table (4) - {e}')

    # Склейка разорванных строк
    new_table_contents = []
    cable_numbers = []
    try:
        for row in table_contents:
            if row[0] not in cable_numbers:
                cable_numbers.append(row[0])
        for cable_number in cable_numbers:
            indices = [j for j, r in enumerate(table_contents) if r[0] == cable_number]
            current_row = table_contents[indices[0]]
            if len(indices) > 1:
                for idx in indices[1:]:
                    for cell_idx in range(len(current_row) - 1):
                        if not is_subset_with_lists(table_contents[idx][cell_idx], current_row[cell_idx]):
                            current_row[cell_idx].extend(table_contents[idx][cell_idx])
            new_table_contents.append(current_row)
        table_contents = new_table_contents
    except Exception as e:
        print(f'Ошибка! extract_all_text_from_table (6) - {e}')

    # Удаление дублей внутри строк
    try:
        new_table_contents = []
        for row in table_contents:
            new_table_contents.append(remove_duplicates(row))
        table_contents = new_table_contents
    except Exception as e:
        print(f'Ошибка! extract_all_text_from_table (7) - {e}')

    # Разбор координат в одной ячейке
    try:
        local_axis = re.compile(r'[+-]?\d{1,8}\.?,?\d{0,3}')
        for i_row in range(len(table_contents)):
            for j_col in range(3, len(table_contents[i_row]) - 1):
                if table_contents[i_row][j_col]:
                    last = table_contents[i_row][j_col][-1]
                    parts = last.strip().split()
                    if len(parts) == 3 and all(local_axis.fullmatch(p) for p in parts):
                        table_contents[i_row][j_col][:-1].extend(parts)
    except Exception as e:
        print(f'Ошибка! extract_all_text_from_table (5) - {e}')

    return table_contents

def extract_all_text_from_docx(docx):
    """Извлекает содержимое всех таблиц из документа .docx"""
    if not docx.tables:
        return []
    all_content = []
    for table in docx.tables:
        if table.rows:
            all_content.extend(extract_all_text_from_table(table))

        # =================================
        # Получение ревизии из колонтитула
        # =================================
    revision = ''
    try:
        section = docx.sections[0]                      # Получаем секцию документа
        header = section.header                         # Получаем верхний колонтитул (header)
        if header.tables:                               # Проверяем, есть ли таблица в колонтитуле
            table1 = header.tables[0]                   # Берём первую таблицу
            if  table1.rows:                            # Берём первую строку таблицы
                first_row = table1.rows[0]
                if first_row.cells:                     # Берём последнюю ячейку в строке (верхняя правая)
                    last_cell = first_row.cells[-1]     # последняя ячейка
                    cell_text = last_cell.text.strip()  # Извлекаем текст из ячейки
                    match = re.search(r'([СC])(\d{2})', cell_text) # паттерн: буква С или C, затем две цифры
                    if match:
                        revision = match.group(0).replace('С', 'C')
    except Exception as e:
        print(f"  ⚠️ Ошибка при извлечении кода из колонтитула {docx}: {e}")
    
    if all_content and revision:
        all_content.append('revision_' + revision)
    return all_content