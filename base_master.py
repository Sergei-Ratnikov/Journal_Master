import re
import openpyxl
from openpyxl import load_workbook
from docx import Document
from pathlib import Path
from datetime import date
import win32com.client
import os
import shutil
import json
from utils import convert_doc_to_docx
from utils import cleanCyrFromLat
from utils import is_subset_with_lists
from utils import all_deep_empty
from utils import get_mismatch_indices
from utils import move_file
import config
from openpyxl.utils import get_column_letter
from openpyxl import Workbook
from openpyxl.styles import PatternFill

# ----- блок функций для поиска актуальной версии журнала

def current_version_finder(dir_in, b_name='Cable base ver.'):
    current_version = 1
    local_list_of_bases_vers = []
    xlsx_files = list(Path(dir_in).glob(b_name + '*.xlsx'))
    for file in xlsx_files:
        dot_index = str(file.stem).rfind('.')
        local_list_of_bases_vers.append(int(str(file.stem)[dot_index + 1:]))
    if not local_list_of_bases_vers:
        print(f'Кабельные базы не найдены в указанной папке')
    else:
        current_version = max(local_list_of_bases_vers) + 1
        print(f'Найдена последняя версия базы {b_name}{current_version - 1}, текущая версия - {current_version}')
    return current_version

# ----- блок функций для работы с docx

def convert_numbering_to_text(doc_path):
    abs_path = os.path.abspath(doc_path)
    word = None
    doc = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False
        doc = word.Documents.Open(abs_path)
        doc.ConvertNumbersToText()
        doc.SaveAs(abs_path)
        return Path(abs_path)
    except Exception as e:
        print(f"Ошибка конвертации нумерации: {e}")
        return None
    finally:
        if doc:
            doc.Close()
        if word:
            word.Quit()

# ----- блок функций для извлечения данных из таблиц

def remove_duplicate_pairs(lst):
    if len(lst) % 2 != 0:
        return lst
    result = []
    for i in range(0, len(lst), 2):
        if lst[i] != lst[i + 1]:
            return lst
        result.append(lst[i])
    return result

def remove_duplicates(lst):
    if not lst:
        return []
    result = []
    skip = 0
    for i in range(len(lst)):
        if skip > 0:
            skip -= 1
            continue
        if i + 2 < len(lst) and lst[i] == ['0'] and lst[i + 1] == ['0'] and lst[i + 2] == ['0']:
            result.extend([['0'], ['0'], ['0']])
            skip = 2
            continue
        if i + 2 < len(lst) and lst[i] == ['-'] and lst[i + 1] == ['-'] and lst[i + 2] == ['-']:
            result.extend([['-'], ['-'], ['-']])
            skip = 2
            continue
        result.append(lst[i])
        j = i + 1
        while j < len(lst) and lst[j] == lst[i]:
            j += 1
        skip = j - i - 1
    return result

def find_first_and_last_sublist_index(big_list, small_list):
    len_small = len(small_list)
    len_big = len(big_list)
    first_index = -1
    last_index = -1
    for i in range(len_big - len_small + 1):
        if big_list[i:i + len_small] == small_list:
            if first_index == -1:
                first_index = i
            last_index = i
    return first_index, last_index

def extract_all_text_from_cell(cell):
    result = []
    if cell:
        for para in cell.paragraphs:
            if para.text.strip():
                text = para.text.strip()
                text = ' '.join(text.split())
                text = text.replace('\xad', '-').replace('\n', ' ')
                result.append(text)
        for inner_table in cell.tables:
            for it_row in inner_table.rows:
                for nested_cell in it_row.cells:
                    inner_table_result = extract_all_text_from_cell(nested_cell)
                    result.extend(inner_table_result)
    return result

def extract_all_text_from_table(table):
    if not table.rows:
        return []
    table_contents = []
    try:
        for row in table.rows:
            current_row = []
            for cell in row.cells:
                current_row.append(extract_all_text_from_cell(cell))
            if len(current_row) > 6 and not all_deep_empty(current_row):
                current_row = remove_duplicate_pairs(current_row)
                table_contents.append(current_row)
    except Exception as e:
        print(f'Ошибка! extract_all_text_from_table (0) - {e}')

    if not table_contents:
        return []

    # Чистка пустых столбцов
    new_table_contents = []
    try:
        for row in table_contents:
            current_row = row
            if all_deep_empty(row[0]) and config.regular_num.search(' '.join(row[1])):
                current_row = row[1:]
            if all_deep_empty(row[-1]) and (any(config.regular_letter_minus.search(s) for s in row[-2]) or all_deep_empty(row[-2])):
                current_row = current_row[:-1]
            new_table_contents.append(current_row)
        table_contents = new_table_contents
    except Exception as e:
        print(f'Ошибка! extract_all_text_from_table (1) - {e}')

    # Удаление строк без номера
    try:
        list_for_delete = []
        for index, row in enumerate(table_contents):
            if row and row[0] and row[0][0]:
                match = config.regular_num.search(row[0][0])
                if match:
                    table_contents[index][0] = [match.group()]
                else:
                    list_for_delete.append(index)
            else:
                list_for_delete.append(index)
        table_contents = [item for idx, item in enumerate(table_contents) if idx not in list_for_delete]
    except Exception as e:
        print(f'Ошибка! extract_all_text_from_table (3) - {e}')

    # Удаление дублей из-за вложенных таблиц
    try:
        list_for_delete = []
        for i in range(len(table_contents) - 1):
            if (len(table_contents[i]) == len(table_contents[i + 1])):
                mismatch = get_mismatch_indices(table_contents[i], table_contents[i + 1])
                if mismatch == [2]:
                    list_for_delete.append(i + 1)
        table_contents = [item for idx, item in enumerate(table_contents) if idx not in list_for_delete]
    except Exception as e:
        print(f'Ошибка! extract_all_text_from_table (4) - {e}')

    # Склейка разорванных строк
    new_table_contents = []
    cable_numbers = []
    try:
        for row in table_contents:
            if row[0] not in cable_numbers:
                cable_numbers.append(row[0])
        for cable_number in cable_numbers:
            indices = [j for j, r in enumerate(table_contents) if r[0] == cable_number]
            current_row = table_contents[indices[0]]
            if len(indices) > 1:
                for idx in indices[1:]:
                    for cell_idx in range(len(current_row) - 1):
                        if not is_subset_with_lists(table_contents[idx][cell_idx], current_row[cell_idx]):
                            current_row[cell_idx].extend(table_contents[idx][cell_idx])
            new_table_contents.append(current_row)
        table_contents = new_table_contents
    except Exception as e:
        print(f'Ошибка! extract_all_text_from_table (6) - {e}')

    # Удаление дублей внутри строк
    try:
        new_table_contents = []
        for row in table_contents:
            new_table_contents.append(remove_duplicates(row))
        table_contents = new_table_contents
    except Exception as e:
        print(f'Ошибка! extract_all_text_from_table (7) - {e}')

    # Разбор координат в одной ячейке
    try:
        local_axis = re.compile(r'[+-]?\d{1,8}\.?,?\d{0,3}')
        for i_row in range(len(table_contents)):
            for j_col in range(3, len(table_contents[i_row]) - 1):
                if table_contents[i_row][j_col]:
                    last = table_contents[i_row][j_col][-1]
                    parts = last.strip().split()
                    if len(parts) == 3 and all(local_axis.fullmatch(p) for p in parts):
                        table_contents[i_row][j_col][:-1].extend(parts)
    except Exception as e:
        print(f'Ошибка! extract_all_text_from_table (5) - {e}')

    return table_contents

def extract_all_text_from_docx(docx):
    if not docx.tables:
        return []
    all_content = []
    for table in docx.tables:
        if table.rows:
            all_content.extend(extract_all_text_from_table(table))
    return all_content

def take_all_docx_from_dir(journals_directory):
    '''
    получение списка всех .docx файлов из указанной директории
    все .doc файлы преобразуются в .docx и перемещаются в отдельную директорию

    '''
    source_dir = Path(journals_directory).resolve()

    docx_files = list(source_dir.glob('*.docx')) 
    doc_files = list(source_dir.glob('*.doc'))
    existing_stems = {f.stem for f in docx_files} # промежуточный список .docx
    files = docx_files.copy()
    for doc_file in doc_files:
        if doc_file.stem not in existing_stems:
            new_path = convert_doc_to_docx(str(doc_file))
            files.append(Path(new_path))
        try:
            doc_backup_dir = source_dir / 'doc files'
            doc_backup_dir.mkdir(exist_ok=True) # создание директории для перемещения .doc файлов
            shutil.move(str(doc_file), str(doc_backup_dir / doc_file.name))
            # print(f"  📁 .doc файл перемещён: {doc_file.name} -> doc files/")
        except Exception as e:
            print(f"  ⚠️ Не удалось переместить {doc_file.name}: {e}")
    return files

# ----- блок функций для работы с KKS и координатами

def parse_coordinate_pair(x_str, y_str):
    x_clean = x_str.strip().replace(',', '.')
    y_clean = y_str.strip().replace(',', '.')
    try:
        x = float(x_clean)
        y = float(y_clean)
        if (4001100 < x < 4002200) and (547700 < y < 549200):
            return (x - 4000000), (y - 547000), 1
    except:
        pass
    try:
        x = float(x_clean)
        y = float(y_clean)
        if (1100 < x < 2200) and (700 < y < 2200):
            return x, y, 1
    except:
        pass
    try:
        x = float(x_clean.replace('.', ''))
        y = float(y_clean.replace('.', ''))
        if (1100000 < x < 2200000) and (700000 < y < 2200000):
            return x / 1000.0, y / 1000.0, 1000
    except:
        pass
    return None, None, 1

def check_and_swap_axes(x, y, bounds, tolerance=100):
    if not bounds or not bounds.get('is_valid'):
        return x, y
    
    def in_range(val, min_val, max_val):
        return (min_val - tolerance) <= val <= (max_val + tolerance)
    
    x_fits = in_range(x, bounds['x_min'], bounds['x_max'])
    y_fits = in_range(y, bounds['y_min'], bounds['y_max'])
    
    if x_fits and y_fits:
        return x, y
    elif in_range(y, bounds['x_min'], bounds['x_max']) and in_range(x, bounds['y_min'], bounds['y_max']):
        # print("Обнаружены перепутанные координаты: X и Y поменяны местами.")
        return y, x
    return x, y

def format_coordinate(value):
    if value is None:
        return ''
    rounded = round(value, 3)
    if rounded.is_integer():
        return str(int(rounded))
    return str(rounded).rstrip('0').rstrip('.')

def normalize_coordinates(kks_room, coords, building_bounds):
    if not coords or len(coords) != 3:
        return None
    x_raw, y_raw, z_raw = coords[0].strip(), coords[1].strip(), coords[2].strip()
    building_kks = None
    if kks_room and building_bounds:
        match = config.regular_KKS_building.search(kks_room)
        if match:
            building_kks = match.group()
            if building_kks not in building_bounds:
                building_kks = kks_room if kks_room in building_bounds else None
        else:
            if kks_room in building_bounds:
                building_kks = kks_room
    bounds = building_bounds.get(building_kks) if building_kks else None
    x_val, y_val, scale = parse_coordinate_pair(x_raw, y_raw)
    if x_val is None or y_val is None:
        # print(f"Не удалось распознать координаты: X='{x_raw}', Y='{y_raw}'")
        return None
    try:
        z_val = float(z_raw.replace(',', '.'))
        if scale == 1000:
            z_val = z_val / 1000.0
    except:
        z_val = 0.0
    x_final, y_final = check_and_swap_axes(x_val, y_val, bounds)
    return [format_coordinate(x_final), format_coordinate(y_final), format_coordinate(z_val)]

def extract_kks_from_list(list_of_strings):
    list_of_kks = []
    if list_of_strings:
        for st in list_of_strings:
            if config.regular_KKS_any.search(st):
                list_of_kks.append(config.regular_KKS_any.search(st).group().strip())
    return list(set(list_of_kks))

def parse_kks_room_and_equip(list_of_KKS, building_bounds):
    KKS_room = ''
    KKS_equipment = ''
    all_buildings = set(building_bounds.keys()) if building_bounds else set()
    if list_of_KKS:
        list_of_KKS = list(set(list_of_KKS))
        for kks in list_of_KKS:
            if config.regular_KKS_room.search(kks):
                building_match = config.regular_KKS_building.search(kks)
                if building_match:
                    building = building_match.group().strip()
                    if building in all_buildings:
                        KKS_room = kks
                        break
            elif config.regular_KKS_building.search(kks) and not KKS_room:
                building = kks.strip()
                if building in all_buildings:
                    KKS_room = building
        for kks in list_of_KKS:
            if config.regular_KKS_equipment.search(kks) and config.regular_KKS_equipment.search(kks).group().strip() != KKS_room:
                KKS_equipment = config.regular_KKS_equipment.search(kks).group().strip()
                break
    return KKS_room, KKS_equipment

# ----- парсер строки

def row_parser(input_row, building_bounds):
    """
    Универсальный парсер строки таблицы кабельного журнала.
    Автоматически находит координаты и KKS независимо от их расположения.
    
    Args:
        input_row: список из функции extract_all_text_from_dir
                  [имя журнала, [ячейка 1], [ячейка 2], ... [ячейка N]]
        building_bounds: словарь с границами зданий из KKS.xlsx
    
    Returns:
        array_row: список из 20 элементов с данными кабеля
    """
    array_row = [''] * 20
    
    if len(input_row) < 7:
        return []
    
    # ========== НОВЫЙ БЛОК: ОЧИСТКА ОТ СИМВОЛОВ ">" ==========
    # Удаляем символы ">" и пробелы в начале строк
    cleaned_row = []
    for cell in input_row:
        if isinstance(cell, list):
            cleaned_cell = []
            for text in cell:
                if isinstance(text, str):
                    # Удаляем ">" в начале и лишние пробелы
                    text = text.lstrip('>').strip()
                    if text:  # только непустые
                        cleaned_cell.append(text)
                else:
                    cleaned_cell.append(text)
            cleaned_row.append(cleaned_cell)
        else:
            cleaned_row.append(cell)
    input_row = cleaned_row
    # ========== КОНЕЦ БЛОКА ==========


    # 0. Журнал
    array_row[0] = cleanCyrFromLat(input_row[0]).strip()
    
    # Проверяем наличие номера кабеля (обычно во второй ячейке)
    if len(input_row) > 1 and input_row[1]:
        array_row[1] = input_row[1][0].replace(',', '.').strip()
    
    # 2. ККС (обычно в третьей ячейке)
    if len(input_row) > 2 and input_row[2]:
        array_row[2] = cleanCyrFromLat(input_row[2][0]).strip()
    
    # 3. Группа (обычно в четвёртой ячейке)
    if len(input_row) > 3 and input_row[3]:
        array_row[3] = input_row[3][0].strip()
    
    # 4. Марка и 5. Сечение — ищем во всех ячейках после третьей
    for cell in input_row[2:]:
        if cell:
            for line in cell:
                if line:
                    # Марка
                    if not array_row[4]:
                        line_clean = line.replace('/', '').strip().split()[0]
                        if config.regular_cableMark.search(line_clean):
                            array_row[4] = line_clean
                        else:
                            for mark in config.arrayCableMarks:
                                if mark in line:
                                    array_row[4] = line
                                    break
                    
                    # Сечение
                    if not array_row[5]:
                        match = config.regular_cableSection.search(line)
                        if match:
                            array_row[5] = match.group().replace('.', ',').replace('х', 'x').replace('×', 'x')
    
    # 7. Длина — обычно в предпоследней или предпредпоследней ячейке
    for idx in [-1, -2, -3]:
        if abs(idx) <= len(input_row) and input_row[idx]:
            length = input_row[idx][0].replace(',', '.')
            
            # Обработка звездочек
            number_pattern = r'\d*\.?\d+'
            pattern_start = rf'^\*{{1,3}}({number_pattern})$'
            pattern_end = rf'^({number_pattern})\*{{1,3}}$'
            
            match = re.match(pattern_start, length)
            if match:
                length = match.group(1)
            match = re.match(pattern_end, length)
            if match:
                length = match.group(1)
            
            if length and length.replace('.', '').replace('-', '').isdigit():
                array_row[7] = length
                break
    
    # 8. Трасса
    if input_row[-1]:
        trace = ', '.join(input_row[-1])
        trace = trace.replace(';', ',').replace('  ', ' ').replace(',,', ',')
        if trace and trace[-1] == ',':
            trace = trace[:-1]
        if trace and len(trace) > 2:
            array_row[8] = trace
    
    # ========== АВТОМАТИЧЕСКИЙ ПОИСК КООРДИНАТ И KKS ==========
    
    # Собираем все текстовые строки из всех ячеек (кроме первых 4 и последних 2)
    int_list = []
    search_start = min(4, len(input_row) - 2)
    search_end = max(search_start + 1, len(input_row) - 2)
    
    for cell in input_row[search_start:search_end]:
        if cell:
            int_list.extend(cell)
    
    # Если не нашли ничего, пробуем расширить диапазон
    if not int_list:
        for cell in input_row[2:]:
            if cell:
                int_list.extend(cell)
    
    # Поиск координат (три числа подряд)
    list_of_axis_start = []
    list_of_axis_end = []
    list_of_KKS_start = []
    list_of_KKS_end = []
    
    # Ищем все возможные координаты в int_list
    coords_indices = []
    for i, val in enumerate(int_list):
        if val and (config.regular_axis_full.search(str(val)) or str(val) == '-' or str(val) == '0' or str(val) == '+0'):
            coords_indices.append(i)
    
    # Ищем три подряд идущие координаты (начало)
    for i in range(len(coords_indices) - 2):
        idx1, idx2, idx3 = coords_indices[i], coords_indices[i+1], coords_indices[i+2]
        if idx2 == idx1 + 1 and idx3 == idx2 + 1:
            list_of_axis_start = [int_list[idx1], int_list[idx2], int_list[idx3]]
            break
    
    # Ищем три подряд идущие координаты с конца (конец)
    for i in range(len(coords_indices) - 1, 1, -1):
        idx1, idx2, idx3 = coords_indices[i-2], coords_indices[i-1], coords_indices[i]
        if idx2 == idx1 + 1 and idx3 == idx2 + 1 and list_of_axis_start:
            # Убеждаемся, что это не те же координаты
            if [int_list[idx1], int_list[idx2], int_list[idx3]] != list_of_axis_start:
                list_of_axis_end = [int_list[idx1], int_list[idx2], int_list[idx3]]
                break
    
    # Если не нашли по три подряд, ищем просто три координаты
    if not list_of_axis_start and len(coords_indices) >= 3:
        list_of_axis_start = [int_list[coords_indices[0]], int_list[coords_indices[1]], int_list[coords_indices[2]]]
    if not list_of_axis_end and len(coords_indices) >= 6:
        list_of_axis_end = [int_list[coords_indices[3]], int_list[coords_indices[4]], int_list[coords_indices[5]]]
    


    # Поиск KKS во всех ячейках
    all_kks_candidates = []
    for cell in input_row:
        if cell:
            kks_list = extract_kks_from_list(cell)
            all_kks_candidates.extend(kks_list)
    
    #     # ========== ВРЕМЕННАЯ ОТЛАДКА ДЛЯ 8.0005 ==========
    
    # if array_row[1] == '8.0005' and array_row[0] == 'AKU.0120.00UKS.0.PS.MB0001-EMB0001':
    #     print(f"\n=== ОТЛАДКА 8.0005 ===")
    #     print(f"all_kks_candidates: {all_kks_candidates}")
    #     print(f"list_of_axis_start: {list_of_axis_start}")
    #     print(f"list_of_axis_end: {list_of_axis_end}")
    #     print(f"int_list: {int_list}")
    #     print(f"\n=== КОНЕЦ ОТЛАДКИ ===")
    # # ========== КОНЕЦ ОТЛАДКИ ==========


    # Разделяем KKS на начало и конец
    # Если есть координаты, KKS до первой координаты — начало, после — конец
    if list_of_axis_start:
        # Находим индекс первой координаты в int_list
        first_coord_idx = -1
        for i, val in enumerate(int_list):
            if val == list_of_axis_start[0]:
                first_coord_idx = i
                break
        
        if first_coord_idx >= 0:
            # KKS до координат — начало
            for cell in input_row:
                if cell:
                    for val in cell:
                        if val and val in all_kks_candidates:
                            # Проверяем, есть ли этот KKS в int_list до координат
                            try:
                                if int_list.index(val) < first_coord_idx:
                                    if val not in list_of_KKS_start:
                                        list_of_KKS_start.append(val)
                            except ValueError:
                                pass
            
            # KKS после координат — конец
            for cell in input_row:
                if cell:
                    for val in cell:
                        if val and val in all_kks_candidates:
                            try:
                                if int_list.index(val) > first_coord_idx + 2:
                                    if val not in list_of_KKS_end:
                                        list_of_KKS_end.append(val)
                            except ValueError:
                                pass
    
    # Если KKS не нашлись по позициям, берем все из соответствующих ячеек
    if not list_of_KKS_start and list_of_axis_start:
        # Берем KKS из ячеек до координат
        for cell in input_row[:search_start+2]:
            if cell:
                kks = extract_kks_from_list(cell)
                list_of_KKS_start.extend(kks)
    
    if not list_of_KKS_end and list_of_axis_end:
        # Берем KKS из ячеек после координат
        for cell in input_row[search_end-2:]:
            if cell:
                kks = extract_kks_from_list(cell)
                list_of_KKS_end.extend(kks)
    
    # Удаляем дубликаты
    list_of_KKS_start = list(set(list_of_KKS_start))
    list_of_KKS_end = list(set(list_of_KKS_end))
    
    # ВЫВОД ДАННЫХ
    # 19. Резервирование
    for cell in input_row[3:5]:
        if cell:
            for line in cell:
                if re.search(r'резерв', line, re.IGNORECASE):
                    array_row[19] = 'Резерв'
                    break
    
    # Определяем KKS помещения и оборудования для начала и конца
    rs, es = parse_kks_room_and_equip(list_of_KKS_start, building_bounds)
    rr, ee = parse_kks_room_and_equip(list_of_KKS_end, building_bounds)
    

    # ========== БЛОК ==========
    # Исправление для кабелей в одном помещении
    if not rr and rs:
        # Если для конца не нашли помещение, но есть координаты
        if len(list_of_axis_end) == 3:
            # Пробуем найти помещение среди start KKS (кроме того, что уже использовано)
            for kks in list_of_KKS_start:
                if kks != rs and config.regular_KKS_room.search(kks):
                    rr = kks
                    break
            # Если всё ещё нет, используем помещение начала
            if not rr:
                rr = rs
                # Помечаем, что это исправление (для отладки)
                # print(f"Исправлено помещение конца: {rr} (было пусто)")
    # Аналогично для обратного случая
    if not rs and rr:
        if len(list_of_axis_start) == 3:
            for kks in list_of_KKS_end:
                if kks != rr and config.regular_KKS_room.search(kks):
                    rs = kks
                    break
            if not rs:
                rs = rr
    # ========== КОНЕЦ БЛОКА ==========

    array_row[9] = rs if rs else ''
    array_row[10] = es if es else ''
    array_row[14] = rr if rr else ''
    array_row[15] = ee if ee else ''
    
    # Нормализуем координаты
    if len(list_of_axis_start) == 3:
        norm_start = normalize_coordinates(rs, list_of_axis_start, building_bounds)
        if norm_start:
            array_row[11], array_row[12], array_row[13] = norm_start[0].replace('.', ','), norm_start[1].replace('.', ','), norm_start[2].replace('.', ',')
        else:
            array_row[11] = list_of_axis_start[0].replace('+', '').strip().replace('.', ',')
            array_row[12] = list_of_axis_start[1].replace('+', '').strip().replace('.', ',')
            array_row[13] = list_of_axis_start[2].replace('+', '').strip().replace('.', ',')
    
    if len(list_of_axis_end) == 3:
        norm_end = normalize_coordinates(rr, list_of_axis_end, building_bounds)
        if norm_end:
            array_row[16], array_row[17], array_row[18] = norm_end[0].replace('.', ','), norm_end[1].replace('.', ','), norm_end[2].replace('.', ',')
        else:
            array_row[16] = list_of_axis_end[0].replace('+', '').strip().replace('.', ',')
            array_row[17] = list_of_axis_end[1].replace('+', '').strip().replace('.', ',')
            array_row[18] = list_of_axis_end[2].replace('+', '').strip().replace('.', ',')
    

    # Специальная обработка для кабелей с координатами только с одной стороны
    # Если есть координаты куда, но нет помещения/оборудования куда
    if (array_row[16] or array_row[17] or array_row[18]) and not (array_row[14] or array_row[15]):
        # Копируем из начала
        if array_row[9] or array_row[10]:
            array_row[14] = array_row[9]   # помещение куда = помещение откуда
            array_row[15] = array_row[10]  # оборудование куда = оборудование откуда

    # Если есть координаты откуда, но нет помещения/оборудования откуда
    if (array_row[11] or array_row[12] or array_row[13]) and not (array_row[9] or array_row[10]):
        # Копируем из конца
        if array_row[14] or array_row[15]:
            array_row[9] = array_row[14]   # помещение откуда = помещение куда
            array_row[10] = array_row[15]  # оборудование откуда = оборудование куда

    # ========== КОНЕЦ ВСТАВКИ ==========


    return array_row

# ----- основная функция

def base_master_start(dir_journals, dir_in):
    '''
    Функция 
    
    Args:
        dir_journals: папка с журналами
        dir_in: папка с кабельной базой, в которую будет сохранена новая база
        все ккс в JSON
    
    Returns:
        кабельная база
    '''

    print("\nЗапуск base_master_start...")

    # Получение списка границ координат зданий в формате
    #     "32UGU": {
    #     "is_valid": true,
    #     "original": "10E; 16N",
    #     "x_min": 1600,
    #     "x_max": 1700,
    #     "y_min": 1000,
    #     "y_max": 1100
    #   }

    with open('KKS_building_bounds.json', 'r', encoding='utf-8') as f:
        building_bounds = json.load(f)

# === поиск старой версии базы
    b_name = 'Cable base ver.'  # шаблон имени базы
    current_version = current_version_finder(dir_in, b_name)
    journals = take_all_docx_from_dir(dir_journals)
    journals_names = [j.stem for j in journals] if journals else []

# === создание новой базы
    wb = openpyxl.Workbook()
    sheetWrite = wb.active
    sheetWrite.title = f"База данных вер. {current_version}"

    heads = [
        [0, 'Журнал', 40],
        [1, 'Номер кабеля', 9],
        [2, 'ККС', 21],
        [3, 'Группа', 7],
        [4, 'Марка', 22],
        [5, 'Сечение', 10],
        [6, 'Диаметр', 8],
        [7, 'Длина', 6],
        [8, 'Трасса', 20],
        [9, 'Откуда помещение', 14],
        [10, 'Откуда оборудование', 20],
        [11, 'Откуда x', 11],
        [12, 'Откуда y', 11],
        [13, 'Откуда z', 11],
        [14, 'Куда помещение', 14],
        [15, 'Куда оборудование', 20],
        [16, 'Куда x', 11],
        [17, 'Куда y', 11],
        [18, 'Куда z', 11],
        [19, 'Резервирование', 8],
        [20, 'Дата добавления в базу', 10],
        [21, 'Версия базы', 7],
        [22, 'id кабеля', 8],
        [23, 'Номер ревизии', 6],
        [24, 'Дата ревизии', 10],
        [25, 'Класс безопасности', 7],
        [26, 'Статус прокладки', 10],
        [27, 'Источник информации ВК/СУПИР', 8],
        [28, 'Статус объединения', 10],
        [29, 'Raw', 25],
        [30, 'min длина', 10],
        [31, 'Длина меньше min', 10]
    ]
    for head in heads:
        col = head[0] + 1
        sheetWrite.cell(row=1, column=col, value=f"{head[0]+1}. {head[1]}")
        sheetWrite.column_dimensions[get_column_letter(col)].width = head[2]
    sheetWrite.freeze_panes = 'A2'

    today = date.today()
    date_string = today.strftime("%Y.%m.%d")
    
    rowWrite = 2
    # Копирование из старой версии базы
    if current_version > 1:
        old_path = Path(dir_in) / f"{b_name}{current_version - 1}.xlsx"
        if old_path.exists():
            rb = load_workbook(old_path, data_only=True)
            sheetRead = rb.active

            rows_to_copy = []
            # составляется список строк старой базы, которые переносятся в новую.
            # переносятся только журналы, которых нет среди журналов в указанной директории dir_journals
            for idx, row in enumerate(sheetRead.iter_rows(min_row=2, values_only=True), start=2):
                journal_name = str(row[0]) if row[0] else ''
                if journal_name and journal_name not in journals_names:
                    rows_to_copy.append(row)

            # графический статус-бар копирования    
            from tqdm import tqdm
            print(f"\n📦 Перенос данных из предыдущей версии ({old_path.name})...")
            with tqdm(total=len(rows_to_copy), desc="Копирование", unit="строка") as pbar:
                for row_data in rows_to_copy:
                    for col_idx, value in enumerate(row_data, start=1):
                        sheetWrite.cell(row=rowWrite, column=col_idx, value=value)
                    rowWrite += 1
                    pbar.update(1)
# проверил
  
# === Обработка журналов
    if journals:
        from tqdm import tqdm
        print("\nОбработка журналов...")
        for journal in tqdm(journals, desc="Обработка", unit="журнал"):
            try:
                j_raw_content = extract_all_text_from_docx(Document(journal))
                if not j_raw_content:
                    continue
                for raw_row in j_raw_content:
                    if len(raw_row) > 6:
                        raw_row.insert(0, journal.stem)
                        processed_row = row_parser(raw_row, building_bounds)

                    if processed_row:
                        # Добавляем дату и версию
                        processed_row.append(date_string)
                        processed_row.append(current_version)
                        
                        # Записываем основные данные
                        for idx, val in enumerate(processed_row, start=1):
                            sheetWrite.cell(row=rowWrite, column=idx, value=val)
                        
                        # Сырая строка (Raw) — в столбец AD (индекс 30, если считать с 0)
                        sheetWrite.cell(row=rowWrite, column=30, value=str(raw_row))
                        
                        # Расчёт минимальной длины кабеля
                        try:
                            length = float(processed_row[7])
                            xyz_from = [float(processed_row[11]), float(processed_row[12]), float(processed_row[13])]
                            xyz_to = [float(processed_row[16]), float(processed_row[17]), float(processed_row[18])]
                            min_len = round(abs(xyz_from[0] - xyz_to[0]) + abs(xyz_from[1] - xyz_to[1]) + abs(xyz_from[2] - xyz_to[2]))
                            sheetWrite.cell(row=rowWrite, column=31, value=min_len)
                            if length < min_len:
                                sheetWrite.cell(row=rowWrite, column=32, value='ДА')
                        except Exception as e:
                            # Если что-то пошло не так — просто пропускаем
                            pass
                        
                        rowWrite += 1


            except Exception as e:
                print(f"Ошибка обработки {journal.stem}: {e}")
    sheetWrite.auto_filter.ref = sheetWrite.dimensions
    output_path = Path(dir_in) / f"{b_name}{current_version}.xlsx"
    wb.save(output_path)
    print(f"\n✅ База данных сохранена: {output_path}")
    print("🏁 Base done!")